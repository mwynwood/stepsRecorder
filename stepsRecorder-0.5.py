"""
Steps Recorder — MS Steps Recorder Replacement (tkinter GUI)

Requirements:
    pip install pynput pillow python-docx

Usage:
    python stepsRecorder.py
"""

import base64
import io
import os
import sys
import threading
import time
from datetime import datetime

try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox
    from pynput import mouse, keyboard
    from PIL import Image, ImageDraw, ImageGrab, ImageTk
except ImportError:
    print("Missing dependencies. Please run:")
    print("    pip install pynput pillow python-docx")
    sys.exit(1)

# ── Config ────────────────────────────────────────────────────────────────────
HIGHLIGHT_RADIUS   = 22
HIGHLIGHT_COLOR    = (255, 60, 60, 190)
HIGHLIGHT_BORDER   = (255, 255, 255, 230)
HIGHLIGHT_BORDER_W = 3
SCREENSHOT_DELAY   = 0.10
DEBOUNCE_MS        = 500
TOGGLE_KEY         = keyboard.Key.f9
OUTPUT_DIR         = "."
THUMB_W, THUMB_H   = 160, 90   # thumbnail size in log
# ─────────────────────────────────────────────────────────────────────────────

steps        = []
recording    = False
_mouse_listener = None
_kb_listener    = None
_lock           = threading.Lock()
_last_click_time = 0.0
_app            = None


# ── Screenshot ────────────────────────────────────────────────────────────────

def take_screenshot(x, y):
    time.sleep(SCREENSHOT_DELAY)
    try:
        shot = ImageGrab.grab()
    except Exception as e:
        log(f"Screenshot failed: {e}")
        return "", None

    img = shot.convert("RGBA")
    overlay = Image.new("RGBA", img.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)
    r, bw = HIGHLIGHT_RADIUS, HIGHLIGHT_BORDER_W
    draw.ellipse([x-r-bw, y-r-bw, x+r+bw, y+r+bw], fill=HIGHLIGHT_BORDER)
    draw.ellipse([x-r, y-r, x+r, y+r], fill=HIGHLIGHT_COLOR)
    draw.line([x-r+5, y, x+r-5, y], fill=(255,255,255,220), width=2)
    draw.line([x, y-r+5, x, y+r-5], fill=(255,255,255,220), width=2)
    result = Image.alpha_composite(img, overlay).convert("RGB")

    # Full-res base64
    buf = io.BytesIO()
    result.save(buf, format="PNG", optimize=True)
    b64 = base64.b64encode(buf.getvalue()).decode()

    # Thumbnail for log
    thumb = result.copy()
    thumb.thumbnail((THUMB_W, THUMB_H), Image.LANCZOS)

    return b64, thumb


# ── Click handler ─────────────────────────────────────────────────────────────

def on_click(x, y, button, pressed):
    global steps, recording, _last_click_time
    if not recording or not pressed:
        return

    now = time.time()
    if (now - _last_click_time) * 1000 < DEBOUNCE_MS:
        return
    _last_click_time = now

    step_num = len(steps) + 1
    btn_name = "Right-click" if button == mouse.Button.right else "Click"

    b64, thumb = take_screenshot(int(x), int(y))

    with _lock:
        steps.append({
            "num":       step_num,
            "action":    btn_name,
            "x":         int(x),
            "y":         int(y),
            "timestamp": datetime.now().strftime("%H:%M:%S"),
            "image_b64": b64,
            "thumb":     thumb,
            "title":     "",
        })

    log(f"Step {step_num}: {btn_name} at ({int(x)}, {int(y)})", thumb=thumb, step_num=step_num)
    if _app:
        _app.after(0, _app.update_count)


# ── Keyboard handler ──────────────────────────────────────────────────────────

def on_key(key):
    if key == TOGGLE_KEY and _app:
        _app.after(0, _app.toggle_recording)


def log(msg, thumb=None, step_num=None):
    if _app:
        _app.after(0, lambda: _app.append_log(msg, thumb=thumb, step_num=step_num))


# ── HTML ──────────────────────────────────────────────────────────────────────

def build_html(steps, title):
    date_str = datetime.now().strftime("%A, %d %B %Y  %H:%M:%S")

    # Table of contents
    toc_items = ""
    for s in steps:
        label = s.get("title") or f'{s["action"]} at ({s["x"]}, {s["y"]})'
        toc_items += f'<li><a href="#step-{s["num"]}"><span class="toc-num">Step {s["num"]}</span> {label}</a></li>\n'

    toc = f"""
    <nav class="toc">
      <div class="toc-title">Contents</div>
      <ol>{toc_items}</ol>
    </nav>""" if steps else ""

    # Step cards
    cards = ""
    for s in steps:
        img_html = (
            f'<img src="data:image/png;base64,{s["image_b64"]}" alt="Step {s["num"]}" loading="lazy">'
            if s["image_b64"] else '<div class="no-img">Screenshot unavailable</div>'
        )
        step_label = s.get("title") or f'{s["action"]} at ({s["x"]}, {s["y"]})'
        cards += f"""
    <div class="card" draggable="true" id="step-{s['num']}" data-step="{s['num']}">
      <div class="card-head">
        <span class="drag-handle" title="Drag to reorder">⠿</span>
        <span class="badge">Step {s['num']}</span>
        <span class="action">{step_label}</span>
        <span class="ts">⏱ {s['timestamp']}</span>
        <button class="delete-btn" title="Delete this step" onclick="deleteStep(this)">✕</button>
      </div>
      <div class="card-body">{img_html}</div>
      <div class="card-desc">
        <input class="title-input" placeholder="Step title (optional)" value="{s.get('title','')}" oninput="updateToc()">
        <textarea class="desc-input" placeholder="Add a caption for this step…" rows="2"></textarea>
      </div>
    </div>"""

    empty = '<div class="empty"><h2>No steps were recorded.</h2></div>' if not steps else ""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{title}</title>
<style>
*,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,sans-serif;background:#0f1117;color:#e2e8f0;min-height:100vh}}

/* Header */
.hdr{{background:linear-gradient(135deg,#1a1d2e,#16213e,#0f3460);border-bottom:1px solid #2d3748;padding:28px 40px;position:sticky;top:0;z-index:100;box-shadow:0 4px 24px rgba(0,0,0,.4)}}
.hdr-inner{{max-width:1100px;margin:0 auto;display:flex;align-items:center;justify-content:space-between;gap:20px;flex-wrap:wrap}}
.hdr h1{{font-size:1.5rem;font-weight:700}}
#report-title{{background:linear-gradient(90deg,#63b3ed,#a78bfa);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;border-bottom:1px dashed #4a5568;padding-bottom:1px;outline:none;cursor:text;transition:border-color .2s}}
#report-title:hover{{border-color:#63b3ed}}
#report-title:focus{{border-color:#a78bfa}}
.hdr p{{font-size:.8rem;color:#718096;margin-top:4px}}
.hdr-right{{display:flex;align-items:center;gap:16px;flex-wrap:wrap}}
.stat-val{{font-size:2rem;font-weight:700;color:#63b3ed;line-height:1}}
.stat-lbl{{font-size:.7rem;color:#718096;text-transform:uppercase;letter-spacing:.08em;margin-top:2px;text-align:center}}
.btn{{border:none;padding:10px 18px;border-radius:8px;font-size:.82rem;font-weight:600;cursor:pointer;transition:opacity .2s;white-space:nowrap;color:#fff}}
.btn:hover{{opacity:.85}}
.btn:disabled{{opacity:.4;cursor:not-allowed}}
.btn-save{{background:linear-gradient(135deg,#3182ce,#7c3aed)}}
.btn-pdf{{background:linear-gradient(135deg,#2f855a,#276749)}}
.btn-group{{display:flex;flex-direction:column;align-items:center;gap:4px}}
.save-notice{{font-size:.72rem;color:#68d391;min-height:1em;text-align:center}}

/* TOC */
.toc{{background:#1a1d2e;border:1px solid #2d3748;border-radius:12px;padding:20px 24px;margin-bottom:32px;max-width:1100px}}
.toc-title{{font-size:.72rem;font-weight:700;text-transform:uppercase;letter-spacing:.1em;color:#718096;margin-bottom:12px}}
.toc ol{{list-style:none;padding:0;columns:2;column-gap:32px}}
.toc li{{margin-bottom:6px;break-inside:avoid}}
.toc a{{color:#63b3ed;text-decoration:none;font-size:.85rem;display:flex;align-items:baseline;gap:8px;transition:color .15s}}
.toc a:hover{{color:#a78bfa}}
.toc-num{{font-size:.72rem;font-weight:700;color:#4a5568;white-space:nowrap}}

/* Layout */
main{{max-width:1100px;margin:0 auto;padding:36px 24px 80px}}

/* Cards */
.card{{background:#1a1d2e;border:1px solid #2d3748;border-radius:12px;overflow:hidden;margin-bottom:24px;transition:border-color .2s,opacity .2s,box-shadow .2s}}
.card:hover{{border-color:#4a5568}}
.card.dragging{{opacity:.4;box-shadow:0 0 0 2px #63b3ed}}
.card.drag-over{{border-color:#a78bfa;box-shadow:0 -3px 0 0 #a78bfa}}
.card-head{{display:flex;align-items:center;gap:10px;padding:13px 18px;background:#16213e;border-bottom:1px solid #2d3748;flex-wrap:wrap}}
.drag-handle{{font-size:1.1rem;color:#4a5568;cursor:grab;padding:0 4px;user-select:none;line-height:1}}
.drag-handle:active{{cursor:grabbing}}
.badge{{background:linear-gradient(135deg,#3182ce,#7c3aed);color:#fff;font-size:.72rem;font-weight:700;padding:3px 10px;border-radius:20px;letter-spacing:.05em;white-space:nowrap}}
.action{{font-size:.86rem;font-weight:600;color:#e2e8f0;flex:1;min-width:0;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}}
.ts{{font-size:.76rem;color:#718096;white-space:nowrap}}
.delete-btn{{background:none;border:none;color:#fc8181;font-size:1rem;cursor:pointer;padding:2px 6px;border-radius:4px;line-height:1;transition:background .15s;margin-left:4px}}
.delete-btn:hover{{background:rgba(252,129,129,.15)}}
.card-body{{padding:14px;background:#0f1117}}
.card-body img{{width:100%;border-radius:6px;border:1px solid #2d3748;display:block;cursor:zoom-in;transition:opacity .2s}}
.card-body img:hover{{opacity:.9}}
.no-img{{padding:40px;text-align:center;color:#4a5568;font-size:.85rem}}
.card-desc{{padding:12px 14px 14px;background:#1a1d2e;border-top:1px solid #2d3748;display:flex;flex-direction:column;gap:8px}}
.title-input{{width:100%;background:#0f1117;border:1px solid #2d3748;border-radius:6px;color:#e2e8f0;font-size:.84rem;font-family:inherit;padding:7px 12px;transition:border-color .2s}}
.title-input:focus{{outline:none;border-color:#a78bfa}}
.title-input::placeholder{{color:#4a5568}}
.desc-input{{width:100%;background:#0f1117;border:1px solid #2d3748;border-radius:6px;color:#e2e8f0;font-size:.84rem;font-family:inherit;padding:9px 12px;resize:vertical;transition:border-color .2s;line-height:1.5}}
.desc-input:focus{{outline:none;border-color:#63b3ed}}
.desc-input::placeholder{{color:#4a5568}}
.empty{{text-align:center;padding:100px 20px;color:#4a5568}}
.empty h2{{color:#718096}}

/* Lightbox */
.lb{{display:none;position:fixed;inset:0;background:rgba(0,0,0,.93);z-index:999;align-items:center;justify-content:center;cursor:zoom-out}}
.lb.on{{display:flex}}
.lb img{{max-width:95vw;max-height:95vh;border-radius:8px;box-shadow:0 0 60px rgba(0,0,0,.8)}}
.lb-x{{position:fixed;top:18px;right:26px;font-size:2rem;color:#fff;cursor:pointer;opacity:.7;line-height:1}}
.lb-x:hover{{opacity:1}}

footer{{text-align:center;padding:20px;font-size:.72rem;color:#4a5568;border-top:1px solid #1a1d2e}}

/* Print */
@media print {{
  body{{background:#fff;color:#000}}
  .hdr{{position:static;background:#fff;border-bottom:2px solid #000;box-shadow:none;padding:20px}}
  .hdr h1{{font-size:1.4rem}}
  #report-title{{background:none;-webkit-text-fill-color:#000;color:#000;border-bottom:none}}
  .hdr p{{color:#555}}
  .hdr-right,.drag-handle,.delete-btn,.save-notice{{display:none!important}}
  .toc{{background:#f9f9f9;border:1px solid #ddd;break-inside:avoid}}
  .toc-title{{color:#333}}
  .toc a{{color:#000}}
  main{{padding:20px 0}}
  .card{{border:1px solid #ccc;border-radius:6px;margin-bottom:20px;break-inside:avoid;page-break-inside:avoid;box-shadow:none}}
  .card-head{{background:#f0f0f0;border-bottom:1px solid #ccc;padding:10px 14px}}
  .badge{{background:#333;-webkit-print-color-adjust:exact;print-color-adjust:exact}}
  .action,.title-input{{color:#000}}
  .ts{{color:#555}}
  .card-body{{background:#fff;padding:10px}}
  .card-body img{{border:1px solid #ddd}}
  .card-desc{{background:#fafafa;border-top:1px solid #ddd;padding:10px 14px}}
  .title-input,.desc-input{{background:#fafafa;border:none;color:#000;font-size:.85rem;resize:none;width:100%}}
  .lb,.btn{{display:none!important}}
  footer{{color:#999;border-top:1px solid #ddd}}
}}
</style>
</head>
<body>
<header class="hdr">
  <div class="hdr-inner">
    <div>
      <h1>📋 <span id="report-title" contenteditable="true" spellcheck="false" title="Click to edit title">{title}</span></h1>
      <p>Recorded on {date_str}</p>
    </div>
    <div class="hdr-right">
      <div>
        <div class="stat-val" id="step-count">{len(steps)}</div>
        <div class="stat-lbl">Steps</div>
      </div>
      <div class="btn-group">
        <button class="btn btn-save" id="save-btn" onclick="saveReport()">💾 Save Report</button>
        <div class="save-notice" id="save-notice"></div>
      </div>
      <button class="btn btn-pdf" onclick="window.print()">🖨 Export PDF</button>
    </div>
  </div>
</header>

<main id="steps-container">
  {toc}
  {cards or empty}
</main>
<footer>Generated by Steps Recorder &nbsp;·&nbsp; {date_str}</footer>

<div class="lb" id="lb">
  <span class="lb-x" id="lb-x">✕</span>
  <img id="lb-img" src="" alt="">
</div>

<script>
// Lightbox
const lb=document.getElementById('lb'),lbImg=document.getElementById('lb-img');
function bindImgClicks(){{document.querySelectorAll('.card-body img').forEach(i=>{{i.onclick=()=>{{lbImg.src=i.src;lb.classList.add('on')}}}})}}
bindImgClicks();
lb.addEventListener('click',e=>{{if(e.target!==lbImg)lb.classList.remove('on')}});
document.getElementById('lb-x').onclick=()=>lb.classList.remove('on');
document.addEventListener('keydown',e=>{{if(e.key==='Escape')lb.classList.remove('on')}});

// Renumber
function updateCount(){{document.getElementById('step-count').textContent=document.querySelectorAll('.card').length}}
function renumber(){{document.querySelectorAll('.card').forEach((c,i)=>c.querySelector('.badge').textContent='Step '+(i+1));updateCount();updateToc();}}

// TOC live update
function updateToc(){{
  const ol=document.querySelector('.toc ol');
  if(!ol)return;
  ol.innerHTML='';
  document.querySelectorAll('.card').forEach((card,i)=>{{
    const badge=card.querySelector('.badge').textContent;
    const titleVal=card.querySelector('.title-input').value.trim();
    const actionVal=card.querySelector('.action').textContent;
    const label=titleVal||actionVal;
    const id=card.id;
    const li=document.createElement('li');
    li.innerHTML=`<a href="#${{id}}"><span class="toc-num">${{badge}}</span> ${{escapeHtml(label)}}</a>`;
    ol.appendChild(li);
  }});
}}

// Title input → update card header live
document.querySelectorAll('.title-input').forEach(inp=>{{
  inp.addEventListener('input',()=>{{
    const card=inp.closest('.card');
    const action=card.querySelector('.action');
    const ds=card.dataset;
    action.textContent=inp.value.trim()||inp.placeholder;
    updateToc();
  }});
}});

// Delete
function deleteStep(btn){{
  const card=btn.closest('.card');
  card.style.transition='opacity .25s,max-height .35s';
  card.style.opacity='0';card.style.maxHeight=card.offsetHeight+'px';
  setTimeout(()=>{{card.style.maxHeight='0';card.style.marginBottom='0';card.style.overflow='hidden';}},50);
  setTimeout(()=>{{card.remove();renumber();}},380);
}}

// Drag & drop
const container=document.getElementById('steps-container');
let dragSrc=null;
function addDragListeners(card){{
  card.addEventListener('dragstart',e=>{{dragSrc=card;setTimeout(()=>card.classList.add('dragging'),0);e.dataTransfer.effectAllowed='move'}});
  card.addEventListener('dragend',()=>{{card.classList.remove('dragging');document.querySelectorAll('.card').forEach(c=>c.classList.remove('drag-over'));renumber()}});
  card.addEventListener('dragover',e=>{{e.preventDefault();if(card!==dragSrc){{document.querySelectorAll('.card').forEach(c=>c.classList.remove('drag-over'));card.classList.add('drag-over')}}}});
  card.addEventListener('drop',e=>{{e.preventDefault();if(dragSrc&&dragSrc!==card){{const cards=[...container.querySelectorAll('.card')];if(cards.indexOf(dragSrc)<cards.indexOf(card))card.after(dragSrc);else card.before(dragSrc)}}card.classList.remove('drag-over')}});
}}
document.querySelectorAll('.card').forEach(addDragListeners);

// Save
async function saveReport(){{
  const btn=document.getElementById('save-btn');
  const notice=document.getElementById('save-notice');
  btn.disabled=true;
  const titleEl=document.getElementById('report-title');
  titleEl.textContent=titleEl.textContent;
  document.querySelectorAll('.desc-input').forEach(ta=>{{ta.innerHTML=escapeHtml(ta.value)}});
  document.querySelectorAll('.title-input').forEach(inp=>{{inp.setAttribute('value',escapeHtml(inp.value))}});
  const html='<!DOCTYPE html>\\n'+document.documentElement.outerHTML;
  try{{
    const handle=await window.showSaveFilePicker({{suggestedName:titleEl.textContent.replace(/[^a-z0-9]/gi,'_')+'.html',types:[{{description:'HTML File',accept:{{'text/html':['.html']}}}}]}});
    const w=await handle.createWritable();
    await w.write(new Blob([html],{{type:'text/html'}}));
    await w.close();
    notice.textContent='✔ Saved!';
  }}catch(e){{
    if(e.name!=='AbortError'){{
      const a=document.createElement('a');
      a.href=URL.createObjectURL(new Blob([html],{{type:'text/html'}}));
      a.download=titleEl.textContent.replace(/[^a-z0-9]/gi,'_')+'_edited.html';
      a.click();
      notice.textContent='✔ Downloaded!';
    }}
  }}
  setTimeout(()=>notice.textContent='',3000);
  btn.disabled=false;
}}

function escapeHtml(s){{return s?s.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;'):''}}
</script>
</body>
</html>"""


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_docx(steps, title, path):
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        messagebox.showerror("Missing library",
            "python-docx is required for Word export.\n\nRun:  pip install python-docx")
        return False

    doc = DocxDocument()

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_para = doc.add_paragraph(f"Recorded on {datetime.now().strftime('%A, %d %B %Y  %H:%M:%S')}")
    date_para.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    date_para.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    for s in steps:
        # Step heading
        step_label = s.get("title") or f'{s["action"]} at ({s["x"]}, {s["y"]})'
        h = doc.add_heading(f'Step {s["num"]}: {step_label}', level=2)

        # Meta line
        meta = doc.add_paragraph(f'⏱ {s["timestamp"]}   •   {s["action"]} at ({s["x"]}, {s["y"]})')
        meta.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        meta.runs[0].font.size = Pt(9)

        # Screenshot
        if s.get("image_b64"):
            img_bytes = base64.b64decode(s["image_b64"])
            img_stream = io.BytesIO(img_bytes)
            try:
                doc.add_picture(img_stream, width=Inches(6))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                doc.add_paragraph("[Screenshot unavailable]")

        # Caption
        if s.get("desc"):
            cap = doc.add_paragraph(s["desc"])
            cap.runs[0].font.italic = True
            cap.runs[0].font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

        doc.add_paragraph()

    doc.save(path)
    return True


# ── tkinter App ───────────────────────────────────────────────────────────────

class App(tk.Tk):
    BG     = "#0f1117"
    CARD   = "#1a1d2e"
    BORDER = "#2d3748"
    FG     = "#e2e8f0"
    MUTED  = "#718096"
    BLUE   = "#3182ce"
    GREEN  = "#38a169"
    RED    = "#e53e3e"
    LOGBG  = "#0d1117"

    def __init__(self):
        super().__init__()
        self.title("Steps Recorder")
        self.resizable(False, False)
        self.configure(bg=self.BG)
        self._mini_win = None
        self._thumb_refs = []   # keep PhotoImage refs alive
        self._blink_job = None
        self._build_ui()
        self.protocol("WM_DELETE_WINDOW", self._on_close)

    # ── Main UI ───────────────────────────────────────────────────────────────

    def _build_ui(self):
        PAD = 16
        BTN = dict(font=("Segoe UI", 10, "bold"), relief="flat",
                   cursor="hand2", padx=14, pady=8, bd=0)

        # Top bar
        top = tk.Frame(self, bg=self.CARD, pady=PAD, padx=PAD)
        top.pack(fill="x")
        tk.Label(top, text="📋 Steps Recorder", font=("Segoe UI", 14, "bold"),
                 bg=self.CARD, fg=self.FG).pack(side="left")
        self._count_var = tk.StringVar(value="0 steps")
        tk.Label(top, textvariable=self._count_var, font=("Segoe UI", 10),
                 bg=self.CARD, fg=self.MUTED).pack(side="right", padx=4)

        # Title field
        mid = tk.Frame(self, bg=self.BG, padx=PAD, pady=10)
        mid.pack(fill="x")
        tk.Label(mid, text="Session title", font=("Segoe UI", 9),
                 bg=self.BG, fg=self.MUTED).pack(anchor="w")
        self._title_var = tk.StringVar(
            value=f"Steps Recording - {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        tk.Entry(mid, textvariable=self._title_var, font=("Segoe UI", 10),
                 bg=self.CARD, fg=self.FG, insertbackground=self.FG,
                 relief="flat", bd=6, width=48).pack(fill="x", pady=(4,0))

        tk.Frame(self, bg=self.BORDER, height=1).pack(fill="x")

        # Buttons
        bf = tk.Frame(self, bg=self.BG, padx=PAD, pady=PAD)
        bf.pack(fill="x")

        self._rec_btn = tk.Button(bf, text="⏺  Start Recording (F9)",
                                  bg=self.BLUE, fg="white",
                                  activebackground="#2b6cb0", activeforeground="white",
                                  command=self.toggle_recording, **BTN)
        self._rec_btn.pack(side="left", padx=(0,6))

        self._finish_btn = tk.Button(bf, text="✔  Finish & Save HTML",
                                     bg=self.GREEN, fg="white",
                                     activebackground="#2f855a", activeforeground="white",
                                     command=lambda: self._finish("html"),
                                     state="disabled", **BTN)
        self._finish_btn.pack(side="left", padx=(0,6))

        self._docx_btn = tk.Button(bf, text="📄  Export Word",
                                   bg="#6b46c1", fg="white",
                                   activebackground="#553c9a", activeforeground="white",
                                   command=lambda: self._finish("docx"),
                                   state="disabled", **BTN)
        self._docx_btn.pack(side="left", padx=(0,6))

        tk.Button(bf, text="✕  Clear", bg="#2d3748", fg=self.FG,
                  activebackground=self.RED, activeforeground="white",
                  command=self._clear, **BTN).pack(side="left", padx=(0,6))

        tk.Button(bf, text="⊡  Mini", bg="#2d3748", fg=self.FG,
                  activebackground="#4a5568", activeforeground="white",
                  command=self._show_mini, **BTN).pack(side="right")

        # Status
        sf = tk.Frame(self, bg=self.BG, padx=PAD)
        sf.pack(fill="x")
        self._dot = tk.Label(sf, text="●", font=("Segoe UI", 11),
                             bg=self.BG, fg=self.BORDER)
        self._dot.pack(side="left")
        self._status_var = tk.StringVar(value="Ready  —  press Start or F9")
        tk.Label(sf, textvariable=self._status_var, font=("Segoe UI", 9),
                 bg=self.BG, fg=self.MUTED).pack(side="left", padx=6)

        tk.Frame(self, bg=self.BORDER, height=1).pack(fill="x", pady=(8,0))

        # Log (canvas + scrollbar for thumbnail support)
        log_outer = tk.Frame(self, bg=self.LOGBG, padx=PAD, pady=PAD)
        log_outer.pack(fill="both", expand=True)

        self._canvas = tk.Canvas(log_outer, bg=self.LOGBG, highlightthickness=0,
                                 width=520, height=320)
        sb = tk.Scrollbar(log_outer, orient="vertical",
                          command=self._canvas.yview)
        self._canvas.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        self._canvas.pack(side="left", fill="both", expand=True)

        self._log_frame = tk.Frame(self._canvas, bg=self.LOGBG)
        self._canvas_win = self._canvas.create_window(
            (0,0), window=self._log_frame, anchor="nw")

        self._log_frame.bind("<Configure>", self._on_log_resize)
        self._canvas.bind("<Configure>", self._on_canvas_resize)

        # Mousewheel
        self._canvas.bind_all("<MouseWheel>",
            lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        tk.Label(self, text="F9 = start/stop  •  Mini mode hides this window",
                 font=("Segoe UI", 8), bg=self.BG, fg=self.BORDER).pack(pady=(0,6))

    def _on_log_resize(self, e):
        self._canvas.configure(scrollregion=self._canvas.bbox("all"))

    def _on_canvas_resize(self, e):
        self._canvas.itemconfig(self._canvas_win, width=e.width)

    # ── Log with thumbnails ───────────────────────────────────────────────────

    def append_log(self, msg, thumb=None, step_num=None):
        row = tk.Frame(self._log_frame, bg=self.LOGBG, pady=4)
        row.pack(fill="x", padx=4)

        if thumb is not None:
            photo = ImageTk.PhotoImage(thumb)
            self._thumb_refs.append(photo)
            lbl = tk.Label(row, image=photo, bg=self.LOGBG,
                           relief="flat", cursor="hand2")
            lbl.pack(side="left", padx=(0,8))

        right = tk.Frame(row, bg=self.LOGBG)
        right.pack(side="left", fill="x", expand=True)

        ts = datetime.now().strftime("%H:%M:%S")
        tk.Label(right, text=f"{ts}  {msg}", font=("Consolas", 9),
                 bg=self.LOGBG, fg=self.MUTED, anchor="w",
                 wraplength=300, justify="left").pack(anchor="w")

        # Separator
        tk.Frame(self._log_frame, bg=self.BORDER, height=1).pack(fill="x", padx=4)

        self._canvas.update_idletasks()
        self._canvas.configure(scrollregion=self._canvas.bbox("all"))
        self._canvas.yview_moveto(1.0)

    # ── Step count ────────────────────────────────────────────────────────────

    def update_count(self):
        n = len(steps)
        s = f"{n} step{'s' if n != 1 else ''}"
        self._count_var.set(s)
        state = "normal" if n > 0 else "disabled"
        self._finish_btn.configure(state=state)
        self._docx_btn.configure(state=state)
        if self._mini_win and self._mini_win.winfo_exists():
            self._mini_win.update_count(n)

    # ── Recording toggle ──────────────────────────────────────────────────────

    def toggle_recording(self):
        global recording, _mouse_listener
        if not recording:
            recording = True
            _mouse_listener = mouse.Listener(on_click=on_click)
            _mouse_listener.start()
            self._rec_btn.configure(text="⏹  Stop Recording (F9)",
                                    bg=self.RED, activebackground="#c53030")
            self._status_var.set("Recording…  click anywhere to capture")
            self._dot.configure(fg="#fc8181")
            self._start_blink()
            self.append_log("Recording started.")
            if self._mini_win and self._mini_win.winfo_exists():
                self._mini_win.set_recording(True)
        else:
            recording = False
            if _mouse_listener:
                _mouse_listener.stop()
            self._rec_btn.configure(text="⏺  Start Recording (F9)",
                                    bg=self.BLUE, activebackground="#2b6cb0")
            self._status_var.set(f"Paused  —  {len(steps)} step(s) captured")
            self._dot.configure(fg=self.BORDER)
            self._stop_blink()
            self.append_log("Recording stopped.")
            if self._mini_win and self._mini_win.winfo_exists():
                self._mini_win.set_recording(False)

    def _start_blink(self):
        self._blink_state = True
        self._blink()

    def _blink(self):
        if not recording:
            return
        self._dot.configure(fg="#fc8181" if self._blink_state else "#7f3d3d")
        self._blink_state = not self._blink_state
        self._blink_job = self.after(600, self._blink)

    def _stop_blink(self):
        if self._blink_job:
            self.after_cancel(self._blink_job)
            self._blink_job = None

    # ── Clear ─────────────────────────────────────────────────────────────────

    def _clear(self):
        global steps
        if steps and not messagebox.askyesno(
                "Clear steps", f"Delete all {len(steps)} captured step(s)?"):
            return
        steps = []
        self._thumb_refs.clear()
        for w in self._log_frame.winfo_children():
            w.destroy()
        self.update_count()
        self.append_log("Cleared.")

    # ── Finish / export ───────────────────────────────────────────────────────

    def _finish(self, fmt):
        if recording:
            self.toggle_recording()
        if not steps:
            messagebox.showwarning("No steps", "No steps have been recorded yet.")
            return

        title = self._title_var.get().strip() or \
                f"Steps Recording - {datetime.now().strftime('%Y-%m-%d %H:%M')}"

        if fmt == "html":
            path = filedialog.asksaveasfilename(
                defaultextension=".html",
                filetypes=[("HTML file", "*.html")],
                initialfile=f"steps_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                title="Save HTML report as…")
            if not path:
                return
            html = build_html(steps, title)
            with open(path, "w", encoding="utf-8") as f:
                f.write(html)
            self.append_log(f"HTML saved: {os.path.basename(path)}")
            messagebox.showinfo("Saved", f"Report saved to:\n{path}")
            import webbrowser
            webbrowser.open(f"file://{os.path.abspath(path)}")

        elif fmt == "docx":
            path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word document", "*.docx")],
                initialfile=f"steps_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                title="Save Word document as…")
            if not path:
                return
            ok = export_docx(steps, title, path)
            if ok:
                self.append_log(f"Word doc saved: {os.path.basename(path)}")
                messagebox.showinfo("Saved", f"Word document saved to:\n{path}")

    # ── Mini mode ─────────────────────────────────────────────────────────────

    def _show_mini(self):
        self.withdraw()
        self._mini_win = MiniWindow(self)
        self._mini_win.set_recording(recording)
        self._mini_win.update_count(len(steps))

    def _restore_from_mini(self):
        if self._mini_win and self._mini_win.winfo_exists():
            self._mini_win.destroy()
        self.deiconify()

    # ── Close ─────────────────────────────────────────────────────────────────

    def _on_close(self):
        global recording, _mouse_listener, _kb_listener
        if recording:
            recording = False
            if _mouse_listener:
                _mouse_listener.stop()
        if _kb_listener:
            _kb_listener.stop()
        self.destroy()


# ── Mini window ───────────────────────────────────────────────────────────────

class MiniWindow(tk.Toplevel):
    def __init__(self, master):
        super().__init__(master)
        self.master = master
        self.title("")
        self.attributes("-topmost", True)
        self.resizable(False, False)
        self.overrideredirect(True)   # borderless
        self.configure(bg="#16213e")
        self._drag_x = 0
        self._drag_y = 0
        self._build()
        # Centre top-right of screen
        sw = self.winfo_screenwidth()
        self.geometry(f"+{sw-220}+40")

    def _build(self):
        outer = tk.Frame(self, bg="#16213e",
                         highlightbackground="#2d3748", highlightthickness=1)
        outer.pack(fill="both", expand=True, padx=1, pady=1)

        # Drag bar
        drag = tk.Frame(outer, bg="#0f3460", cursor="fleur")
        drag.pack(fill="x")
        tk.Label(drag, text="Steps Recorder", font=("Segoe UI", 8, "bold"),
                 bg="#0f3460", fg="#a0aec0").pack(side="left", padx=8, pady=4)
        tk.Button(drag, text="✕", font=("Segoe UI", 8), bg="#0f3460", fg="#fc8181",
                  relief="flat", bd=0, cursor="hand2",
                  command=self.master._restore_from_mini).pack(side="right", padx=4)
        drag.bind("<ButtonPress-1>", self._start_drag)
        drag.bind("<B1-Motion>", self._do_drag)

        # Content row
        row = tk.Frame(outer, bg="#16213e", padx=10, pady=8)
        row.pack(fill="x")

        self._rec_btn = tk.Button(row, text="⏺  Record",
                                  font=("Segoe UI", 9, "bold"),
                                  bg="#3182ce", fg="white",
                                  activebackground="#2b6cb0",
                                  relief="flat", bd=0, padx=10, pady=6,
                                  cursor="hand2",
                                  command=self.master.toggle_recording)
        self._rec_btn.pack(side="left", padx=(0,10))

        self._count_lbl = tk.Label(row, text="0 steps",
                                   font=("Segoe UI", 9),
                                   bg="#16213e", fg="#718096")
        self._count_lbl.pack(side="left")

    def _start_drag(self, e):
        self._drag_x = e.x_root - self.winfo_x()
        self._drag_y = e.y_root - self.winfo_y()

    def _do_drag(self, e):
        self.geometry(f"+{e.x_root - self._drag_x}+{e.y_root - self._drag_y}")

    def set_recording(self, active):
        if active:
            self._rec_btn.configure(text="⏹  Stop", bg="#e53e3e",
                                    activebackground="#c53030")
        else:
            self._rec_btn.configure(text="⏺  Record", bg="#3182ce",
                                    activebackground="#2b6cb0")

    def update_count(self, n):
        self._count_lbl.configure(text=f"{n} step{'s' if n != 1 else ''}")


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    global _app, _kb_listener

    _app = App()

    _kb_listener = keyboard.Listener(on_press=on_key)
    _kb_listener.start()

    _app.mainloop()


if __name__ == "__main__":
    main()
